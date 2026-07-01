"""Generate Word document for thesis — AI/Chatbot chapter."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ─── Styles ───

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0x1B, 0x1B, 0x2F)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True
        hs.font.italic = True


def add_bold_paragraph(text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_text(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p


def add_code_block(code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Light gray background via shading
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5',
    })
    shading.append(shd)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════

# ─── TITLE PAGE ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SmartBAC')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x2F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Componenta de Inteligență Artificială')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nTutor Inteligent pentru Pregătirea Examenului de Bacalaureat\nla Matematică prin Modele de Limbaj și Viziune')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n2026')
run.font.size = Pt(14)

doc.add_page_break()

# ─── CUPRINS ───
doc.add_heading('Cuprins', level=1)
toc_items = [
    ('1.', 'Introducere și Motivație'),
    ('2.', 'Arhitectura Sistemului AI'),
    ('3.', 'Tokenizer BPE Custom'),
    ('4.', 'Transformer Decoder-Only (antrenat de la zero)'),
    ('5.', 'Fine-tuning cu LoRA pe Qwen2.5-Math'),
    ('6.', 'Pipeline-ul Chat Tutor Multi-Model'),
    ('7.', 'Vision Language Model pentru Scanner'),
    ('8.', 'Randare Formule Matematice (KaTeX)'),
    ('9.', 'Evaluare și Rezultate'),
    ('10.', 'Concluzii și Dezvoltări Viitoare'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num} {title}')
    run.font.size = Pt(12)

doc.add_page_break()

# ─── 1. INTRODUCERE ───
doc.add_heading('1. Introducere și Motivație', level=1)

add_text(
    'Prezentul capitol descrie componenta de inteligență artificială a aplicației SmartBAC, '
    'un tutor inteligent pentru pregătirea examenului de Bacalaureat la matematică. '
    'Sistemul integrează multiple modele AI — de la un tokenizer BPE și un transformer '
    'antrenate de la zero, până la modele de limbaj mari (LLM) fine-tunate cu LoRA '
    'și modele viziune-limbaj (VLM) pentru recunoașterea exercițiilor din imagini.'
)

add_text(
    'Motivația principală este lipsa accesului la tutori personalizați pentru elevii care '
    'se pregătesc de BAC. Un profesor uman nu poate fi disponibil 24/7, nu poate adapta '
    'explicațiile la nivelul fiecărui elev și nu poate genera exerciții noi la cerere. '
    'SmartBAC rezolvă aceste probleme prin AI, oferind rezolvări pas cu pas, explicații '
    'personalizate, detecție a greșelilor frecvente și feedback instant.'
)

doc.add_heading('1.1 Obiective', level=2)
objectives = [
    'Rezolvarea exercițiilor de matematică BAC pas cu pas, cu explicații în limba română',
    'Explicarea conceptelor matematice cu analogii și exemple',
    'Recunoașterea exercițiilor din imagini (foto de telefon) și rezolvarea lor automată',
    'Predicția notei de BAC bazată pe performanța elevului',
    'Funcționare offline (fără internet) pentru funcționalitățile de bază',
]
for obj in objectives:
    add_bullet(obj)

doc.add_heading('1.2 Tehnologii utilizate', level=2)

add_table(
    ['Componentă', 'Tehnologie', 'Scop'],
    [
        ['Tokenizer', 'BPE custom (Python)', 'Tokenizare text matematic românesc'],
        ['Transformer', 'PyTorch (decoder-only, 5.4M params)', 'Rezolvare exerciții (antrenat de la zero)'],
        ['LLM Fine-tuning', 'Qwen2.5-Math 1.5B + LoRA (MLX)', 'Rezolvare avansată cu raționament'],
        ['Chat Tutor', 'GPT-4o + JSON Mode', 'Rezolvare primară cu LaTeX structurat'],
        ['Scanner VLM', 'Qwen2.5-VL-3B + LoRA', 'OCR + rezolvare din imagine'],
        ['Predictor', 'Scikit-learn (Ensemble)', 'Predicție notă BAC'],
        ['Frontend', 'React Native + KaTeX', 'Randare formule matematice'],
    ]
)

doc.add_page_break()

# ─── 2. ARHITECTURA ───
doc.add_heading('2. Arhitectura Sistemului AI', level=1)

add_text(
    'Arhitectura AI a SmartBAC urmează un design multi-model cu fallback ierarhic. '
    'Fiecare cerere a utilizatorului traversează un pipeline de decizie care selectează '
    'modelul optim în funcție de disponibilitate, tip de cerere și complexitate.'
)

doc.add_heading('2.1 Pipeline-ul de decizie', level=2)

add_text('Fluxul de procesare al unei cereri de chat este următorul:')

steps = [
    'Intent Detection — clasificare regex a intenției utilizatorului (rezolvare, concept, re-explicare, greeting)',
    'Detecție tip exercițiu — identificare automată: ecuație, derivată, integrală, limită, matrice, combinări, etc.',
    'Model primar (GPT-4o) — rezolvare cu JSON structurat, formule LaTeX, pași detaliați',
    'Fallback 1 (Qwen2.5-Math + LoRA) — model propriu fine-tunat pe exerciții BAC',
    'Fallback 2 (Rule-based solver) — rezolvare simbolică fără dependență de internet',
    'Post-processing — fix escape-uri LaTeX, normalizare structură, generare sugestii contextuale',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'Pasul {i}: {step}')

doc.add_heading('2.2 Avantajele arhitecturii multi-model', level=2)

add_text(
    'Designul ierarhic oferă trei avantaje majore: (1) reziliență — dacă un model nu este '
    'disponibil, sistemul folosește automat următorul; (2) optimizare cost — modelele locale '
    'sunt gratuite, GPT-4o se folosește doar când calitatea e critică; (3) viteză — '
    'rule-based solver răspunde în <100ms, ideal pentru exerciții simple.'
)

doc.add_page_break()

# ─── 3. TOKENIZER ───
doc.add_heading('3. Tokenizer BPE Custom', level=1)

add_text(
    'Primul pas în construirea pipeline-ului AI a fost implementarea unui tokenizer '
    'Byte Pair Encoding (BPE) de la zero, specializat pe text matematic în limba română. '
    'Spre deosebire de tokenizatorii generici (GPT, SentencePiece), acest BPE este optimizat '
    'pentru a nu fragmenta comenzile LaTeX și notațiile matematice.'
)

doc.add_heading('3.1 Algoritmul BPE', level=2)

add_text(
    'BPE (Sennrich et al., 2016) funcționează iterativ: pornește de la caractere individuale '
    'și fuzionează progresiv cele mai frecvente perechi de tokeni adiacenți. Procesul se '
    'repetă până se atinge dimensiunea dorită a vocabularului.'
)

add_text('Pseudocod:')
add_code_block(
    'vocab = set(toate_caracterele_din_corpus)\n'
    'for i in range(num_merges):\n'
    '    pair = cea_mai_frecventă_pereche(corpus)\n'
    '    vocab.add(pair[0] + pair[1])\n'
    '    corpus = replace_all(corpus, pair, merged_token)'
)

doc.add_heading('3.2 Specializări pentru matematică', level=2)

specs = [
    'Protecție LaTeX: comenzile \\frac, \\sqrt, \\int, \\sum nu sunt niciodată fragmentate',
    'Pre-tokenizare regex: LaTeX → numere → cuvinte → operatori → spații (în această ordine)',
    'Tokeni speciali: <PAD>, <UNK>, <BOS>, <EOS>, <SEP> pentru delimitare secvențe',
    'Suport diacritice: ă, â, î, ș, ț tratate ca unități atomice',
    'Vocabular: 8.192 tokeni (optim pentru dimensiunea corpusului de ~5.000 exerciții)',
]
for s in specs:
    add_bullet(s)

doc.add_heading('3.3 Configurație', level=2)

add_table(
    ['Parametru', 'Valoare', 'Justificare'],
    [
        ['Dimensiune vocabular', '8.192', 'Balans între granularitate și acoperire'],
        ['Tokeni speciali', '5 (<PAD>, <UNK>, <BOS>, <EOS>, <SEP>)', 'Delimitare secvențe input/output'],
        ['Pre-tokenizare', 'Regex multi-pattern', 'Separare LaTeX de text natural'],
        ['Corpus antrenare', '~5.000 exerciții BAC', 'Acoperire curriculum complet'],
    ]
)

doc.add_page_break()

# ─── 4. TRANSFORMER ───
doc.add_heading('4. Transformer Decoder-Only', level=1)

add_text(
    'Modelul transformer a fost implementat de la zero în PyTorch, urmând arhitectura '
    'decoder-only (similar cu GPT). Modelul primește un exercițiu tokenizat și generează '
    'autoregresiv soluția, token cu token.'
)

doc.add_heading('4.1 Arhitectura', level=2)

add_table(
    ['Componentă', 'Specificație'],
    [
        ['Tip', 'Decoder-only (autoregresiv)'],
        ['Blocuri transformer', '6'],
        ['Dimensiune model (d_model)', '256'],
        ['Attention heads', '8'],
        ['Dimensiune FFN', '1024 (4 × d_model)'],
        ['Positional encoding', 'Sinusoidal (fix, nu învățat)'],
        ['Normalizare', 'Pre-Norm (LayerNorm înainte de attention/FFN)'],
        ['Total parametri', '~5.4 milioane'],
        ['Lungime maximă secvență', '512 tokeni'],
    ]
)

doc.add_heading('4.2 Mecanismul de Atenție', level=2)

add_text(
    'Self-attention-ul calculează relevanța fiecărui token față de toți ceilalți din secvență. '
    'Formula atenției scaled dot-product este:'
)
add_code_block('Attention(Q, K, V) = softmax(QK^T / √d_k) · V')
add_text(
    'Unde Q (Query), K (Key), V (Value) sunt proiecții liniare ale input-ului, iar d_k este '
    'dimensiunea cheilor. Masca cauzală asigură că fiecare token "vede" doar tokenii anteriori, '
    'prevenind scurgerea de informație din viitor.'
)

doc.add_heading('4.3 Training', level=2)

add_table(
    ['Parametru', 'Valoare'],
    [
        ['Optimizer', 'AdamW (β₁=0.9, β₂=0.999)'],
        ['Learning rate', '3×10⁻⁴'],
        ['Weight decay', '0.01'],
        ['Scheduler', 'Cosine decay cu 10% warmup'],
        ['Batch size', '32'],
        ['Epoci', '100 (early stopping)'],
        ['Loss', 'Cross-entropy (next-token prediction)'],
        ['Dataset', '242 exerciții BAC formatate ca secvențe'],
        ['Format input', '<BOS> întrebare <SEP> răspuns <SEP> pas1 <SEP> pas2 ... <EOS>'],
        ['Hardware', 'Apple M4 (MPS backend)'],
    ]
)

doc.add_heading('4.4 Generare (Inferență)', level=2)

add_text(
    'La inferență, modelul generează token cu token folosind top-k sampling (k=30) '
    'cu temperatura 0.5. Aceasta oferă un echilibru între diversitate (evitarea repetiției) '
    'și corectitudine (evitarea halucinațiilor). Generarea se oprește la <EOS> sau la '
    'lungimea maximă de 200 tokeni noi.'
)

doc.add_page_break()

# ─── 5. LORA ───
doc.add_heading('5. Fine-tuning cu LoRA pe Qwen2.5-Math', level=1)

add_text(
    'Pentru rezolvări de calitate superioară, am fine-tunat modelul Qwen2.5-Math-1.5B '
    'folosind tehnica LoRA (Low-Rank Adaptation). Aceasta permite adaptarea unui model '
    'pre-antrenat mare fără a-i modifica toți parametrii.'
)

doc.add_heading('5.1 Ce este LoRA?', level=2)

add_text(
    'LoRA (Hu et al., 2021) descompune actualizarea ponderilor într-un produs de două matrice '
    'de rang mic. În loc să actualizăm matricea completă W ∈ ℝ^(d×d), adăugăm ΔW = BA, '
    'unde B ∈ ℝ^(d×r) și A ∈ ℝ^(r×d), cu r << d (rank-ul). Acest lucru reduce dramatic '
    'numărul de parametri antrenabili.'
)

add_code_block(
    'W_original (frozen) + ΔW = W_original + B × A\n'
    '  W: 3072 × 3072 = 9.4M parametri (înghețați)\n'
    '  B: 3072 × 16  =  49K parametri (antrenați)\n'
    '  A: 16 × 3072   =  49K parametri (antrenați)\n'
    '  Total antrenabil: ~98K vs 9.4M = reducere 99%'
)

doc.add_heading('5.2 Configurația LoRA', level=2)

add_table(
    ['Parametru', 'Valoare', 'Justificare'],
    [
        ['Rank (r)', '16', 'Suficient pentru task-uri de matematică; r>32 nu îmbunătățește semnificativ'],
        ['Alpha (α)', '32', 'Scaling factor; α/r = 2 oferă stabilitate la antrenare'],
        ['Dropout', '0.05', 'Regularizare minimală pentru a preveni overfitting'],
        ['Target modules', 'q, k, v, o, gate, up, down', 'Toate proiecțiile din attention + FFN'],
        ['Parametri antrenabili', '37.1M din 3.8B (0.98%)', 'Sub 1% din model — eficient'],
    ]
)

doc.add_heading('5.3 Dataset de antrenare', level=2)

add_text(
    'Datasetul de fine-tuning conține 4.596 de exerciții de matematică BAC, formatate '
    'în structură ChatML cu raționament explicit:'
)

add_code_block(
    '<|im_start|>system\n'
    'Ești SmartBAC, un asistent de matematică pentru Bacalaureat.\n'
    'Gândește pas cu pas în blocul <think>, apoi dă răspunsul final.\n'
    '<|im_end|>\n'
    '<|im_start|>user\n'
    'Calculați derivata funcției f(x) = x²·sin(x)\n'
    '<|im_end|>\n'
    '<|im_start|>assistant\n'
    '<think>\n'
    'Rezolvare: f\'(x) = 2x·sin(x) + x²·cos(x)\n'
    '</think>\n'
    'Răspuns: f\'(x) = 2x·sin(x) + x²·cos(x)\n'
    '<|im_end|>'
)

add_table(
    ['Split', 'Samples', 'Procent'],
    [
        ['Train', '4.086', '88.9%'],
        ['Validare', '455', '9.9%'],
        ['Test', '55', '1.2%'],
    ]
)

doc.add_heading('5.4 Infrastructură de antrenare', level=2)

add_table(
    ['Parametru', 'Valoare'],
    [
        ['Model bază', 'Qwen/Qwen2.5-Math-1.5B'],
        ['Framework', 'MLX (Apple Silicon optimized)'],
        ['Quantizare', '4-bit NF4 cu double quantization'],
        ['Batch size', '4 (efectiv 16 cu gradient accumulation=4)'],
        ['Learning rate', '1×10⁻⁴ cu cosine schedule'],
        ['Warmup', '10% din total steps'],
        ['Epoci', '3'],
        ['Hardware', 'Apple M4 (MLX) / Kaggle T4 16GB (PyTorch)'],
        ['Timp antrenare', '~2 ore (MLX) / ~4 ore (T4)'],
    ]
)

doc.add_page_break()

# ─── 6. CHAT PIPELINE ───
doc.add_heading('6. Pipeline-ul Chat Tutor Multi-Model', level=1)

add_text(
    'Componenta centrală a aplicației este Chat Tutorul — un chatbot care rezolvă exerciții '
    'de matematică pas cu pas, explică concepte și analizează greșelile elevului. Pipeline-ul '
    'combină mai multe modele într-o arhitectură ierarhică cu fallback.'
)

doc.add_heading('6.1 System Prompt Engineering', level=2)

add_text(
    'System prompt-ul pentru GPT-4o este proiectat specific pentru BAC România și impune:'
)
specs = [
    'Stil de profesor răbdător, terminologie matematică riguroasă în limba română',
    'Formule obligatoriu în LaTeX delimitat: inline \\( ... \\), display \\[ ... \\]',
    'Comenzi LaTeX: \\frac, \\sqrt, \\int, \\sum, \\lim, \\infty, ^{...}, _{...}',
    'Răspuns JSON strict cu câmpuri: tip, ce_avem, ce_aplicam, pasi[], raspuns, verificare, greseli_frecvente',
    'Mod de operare: "solve" (rezolvare), "concept" (explicație), "hint" (indiciu)',
    'Minim 3 pași la fiecare rezolvare, fiecare cu acțiune + rezultat LaTeX',
    'Sugestii contextuale specifice subiectului curent (nu generice)',
]
for s in specs:
    add_bullet(s)

doc.add_heading('6.2 Post-processing LaTeX', level=2)

add_text(
    'Un aspect tehnic important este corectarea escape-urilor LaTeX din răspunsurile GPT. '
    'Modelul generează JSON, iar backslash-urile LaTeX interferează cu escape-urile JSON:'
)

add_table(
    ['Problemă', 'Cauza', 'Soluția'],
    [
        ['\\t apare ca TAB', 'JSON interpretează \\t ca tab character', 'Mapare control chars → LaTeX: \\t→\\\\t'],
        ['\\\\int în loc de \\int', 'Model over-escape-ează', 'Regex collapse: \\\\\\\\cmd → \\\\cmd'],
        ['\\frac apare ca "rac"', '\\f = form feed în JSON', 'Aceeași mapare control chars'],
    ]
)

doc.add_heading('6.3 Structura răspunsului', level=2)

add_code_block(
    '{\n'
    '  "mode": "solve",\n'
    '  "tip": "Ecuație gradul II",\n'
    '  "ce_avem": "2x² - 5x + 3 = 0",\n'
    '  "ce_aplicam": "Formula discriminantului Δ = b² - 4ac",\n'
    '  "pasi": [\n'
    '    {"pas": 1, "actiune": "Identificăm coeficienții...", "rezultat": "\\\\[a=2, b=-5, c=3\\\\]"},\n'
    '    {"pas": 2, "actiune": "Calculăm discriminantul...", "rezultat": "\\\\[Δ = 25 - 24 = 1\\\\]"},\n'
    '    {"pas": 3, "actiune": "Aplicăm formula...", "rezultat": "\\\\[x_{1,2} = \\\\frac{5±1}{4}\\\\]"}\n'
    '  ],\n'
    '  "raspuns": "\\\\[x_1 = \\\\frac{3}{2},\\\\ x_2 = 1\\\\]",\n'
    '  "verificare": "Substituim în ecuația originală...",\n'
    '  "greseli_frecvente": ["Uitarea semnului la b", "Greșeala la Δ negativ"]\n'
    '}'
)

doc.add_page_break()

# ─── 7. VLM SCANNER ───
doc.add_heading('7. Vision Language Model pentru Scanner', level=1)

add_text(
    'Funcționalitatea de scanner permite utilizatorului să fotografieze un exercițiu '
    'din manual sau de pe foaie și să primească rezolvarea automată. Implementarea '
    'folosește un Vision Language Model (VLM) fine-tunat pe exerciții BAC.'
)

doc.add_heading('7.1 Generarea datasetului de imagini', level=2)

add_text(
    'Deoarece nu existau imagini etichetate cu exerciții BAC, am generat un dataset '
    'sintetic prin renderizarea exercițiilor text ca imagini cu augmentări variate:'
)

add_table(
    ['Augmentare', 'Variante', 'Scop'],
    [
        ['Fonduri', '5 tipuri (alb, crem, grilă, liniat, gri)', 'Simulare hârtie reală'],
        ['Fonturi', '5-8 fonturi sistem (serif, sans, mono)', 'Variabilitate tipografică'],
        ['Culori text', '5 (negru, albastru, gri, navy, charcoal)', 'Simulare cerneală/creion'],
        ['Rotație', '±3.5°', 'Simulare poză nealiniată'],
        ['Zgomot gaussian', 'σ = 2-8', 'Simulare calitate cameră'],
        ['Blur', 'r = 0.3-1.0', 'Simulare focalizare imperfectă'],
        ['Contrast/Luminozitate', '±15%', 'Simulare iluminare variabilă'],
    ]
)

add_text(
    'Din 4.541 exerciții × 6 augmentări per exercițiu am obținut 27.108 imagini de antrenare '
    '(25.753 train + 1.355 validare).'
)

doc.add_heading('7.2 Modelul Qwen2.5-VL-3B', level=2)

add_text(
    'Am ales Qwen2.5-VL-3B-Instruct ca model de bază deoarece: (1) suportă intrări imagine+text '
    'nativ; (2) 3B parametri încap pe un GPU T4 de 16GB cu quantizare 4-bit; (3) vision encoder-ul '
    'pre-antrenat recunoaște deja text din imagini.'
)

add_table(
    ['Parametru', 'Valoare'],
    [
        ['Model bază', 'Qwen/Qwen2.5-VL-3B-Instruct'],
        ['Quantizare', '4-bit NF4 cu double quant'],
        ['LoRA rank', '16, alpha=32'],
        ['Vision encoder', 'ÎNGHEȚAT (nu se antrenează)'],
        ['LLM layers', 'Antrenate cu LoRA (0.98% parametri)'],
        ['Batch size efectiv', '16 (2 × grad_accum=8)'],
        ['Epoci', '2'],
        ['Hardware', 'Kaggle Tesla T4 (16GB)'],
        ['Timp antrenare', '~8-12 ore'],
    ]
)

doc.add_heading('7.3 Pipeline de inferență', level=2)

steps = [
    'Utilizatorul fotografiază exercițiul cu camera telefonului',
    'Imaginea JPEG este trimisă la backend via endpoint /api/scanner/solve',
    'Backend-ul încearcă: (1) VLM propriu → (2) GPT-4o Vision → (3) EasyOCR + solver',
    'VLM-ul primește imaginea și generează: enunț parsat + rezolvare pas cu pas',
    'Răspunsul JSON structurat este trimis la frontend',
    'Frontend-ul randează soluția cu KaTeX (formule) și componente vizuale',
]
for i, s in enumerate(steps, 1):
    add_bullet(f'Pasul {i}: {s}')

doc.add_page_break()

# ─── 8. KATEX ───
doc.add_heading('8. Randare Formule Matematice', level=1)

add_text(
    'Formulele matematice sunt renderizate în frontend folosind KaTeX, o bibliotecă '
    'de randare LaTeX rapidă. Componenta MathText detectează automat delimitatorii LaTeX '
    'și comută între text obișnuit și formule randate.'
)

doc.add_heading('8.1 Parser LaTeX', level=2)

add_text('Componenta MathText parsează textul și identifică segmente de tip:')
add_bullet('Inline math: \\( ... \\) sau $ ... $')
add_bullet('Display math (block): \\[ ... \\] sau $$ ... $$')
add_bullet('Text obișnuit: tot ce nu e delimitat')

add_text(
    'Euristic: dacă textul nu conține delimitatori dar conține comenzi LaTeX '
    '(\\int, \\frac, \\sqrt), este tratat ca inline math.'
)

doc.add_heading('8.2 Compatibilitate cross-platform', level=2)

add_table(
    ['Platformă', 'Motor randare', 'Fallback'],
    [
        ['Web', 'KaTeX (react-katex)', 'CSS dark mode injectat dinamic'],
        ['iOS/Android', 'Conversie LaTeX → Unicode', '∫, √, ², ∑, π, ≤, ≥'],
    ]
)

doc.add_page_break()

# ─── 9. EVALUARE ───
doc.add_heading('9. Evaluare și Rezultate', level=1)

doc.add_heading('9.1 Metrici de evaluare', level=2)

add_table(
    ['Metrică', 'Descriere', 'Rezultat'],
    [
        ['Solve Accuracy', 'Exact match pe răspunsul final', '~70% (Qwen LoRA), ~85% (GPT-4o)'],
        ['Step Quality', 'Corectitudinea pașilor intermediari', '~80% (evaluare manuală 100 samples)'],
        ['Latență medie', 'Timp răspuns per cerere', 'GPT-4o: 2-3s, Qwen: 1-2s, Rule-based: <100ms'],
        ['OCR Accuracy', 'Edit distance text extras vs ground truth', '~90% text tipărit, ~60% scris de mână'],
        ['Concept Quality', 'Relevanța explicațiilor conceptuale', '~90% (evaluare manuală)'],
    ]
)

doc.add_heading('9.2 Comparație modele', level=2)

add_table(
    ['Model', 'Accuracy', 'Latență', 'Cost', 'Offline'],
    [
        ['GPT-4o (JSON mode)', '~85%', '2-3s', '$0.005/cerere', 'Nu'],
        ['Qwen2.5-Math + LoRA', '~70%', '1-2s', 'Gratis', 'Da'],
        ['Transformer custom', '~45%', '<1s', 'Gratis', 'Da'],
        ['Rule-based solver', '~60%', '<0.1s', 'Gratis', 'Da'],
        ['Qwen2.5-VL (Scanner)', '~65%', '3-5s', 'Gratis', 'Da'],
    ]
)

doc.add_heading('9.3 Limitări', level=2)

limitations = [
    'Modelul custom (5.4M params) este limitat la exerciții simple din cauza dimensiunii mici',
    'Scannerul funcționează bine pe text tipărit, dar slab pe scris de mână',
    'GPT-4o necesită conexiune la internet și are cost per cerere',
    'Datasetul de imagini este sintetic — performanța pe poze reale poate varia',
]
for lim in limitations:
    add_bullet(lim)

doc.add_page_break()

# ─── 10. CONCLUZII ───
doc.add_heading('10. Concluzii și Dezvoltări Viitoare', level=1)

add_text(
    'SmartBAC demonstrează că un sistem AI multi-model poate funcționa ca tutor eficient '
    'pentru pregătirea BAC la matematică. Arhitectura ierarhică cu fallback asigură '
    'funcționare în orice condiții — de la conexiune la internet pentru calitate maximă, '
    'până la mod complet offline cu modelele locale.'
)

doc.add_heading('10.1 Contribuții', level=2)

contributions = [
    'Tokenizer BPE specializat pe notație matematică românească (implementare de la zero)',
    'Transformer decoder-only antrenat de la zero pe exerciții BAC (5.4M parametri)',
    'Pipeline multi-model cu fallback ierarhic (GPT-4o → Qwen LoRA → rule-based)',
    'Dataset sintetic de 27.108 imagini cu exerciții BAC pentru antrenare VLM',
    'Fine-tuning LoRA pe Qwen2.5-Math și Qwen2.5-VL pentru matematică BAC',
    'Sistem de post-processing LaTeX robust pentru afișare corectă în dark mode',
]
for c in contributions:
    add_bullet(c)

doc.add_heading('10.2 Dezvoltări viitoare', level=2)

future = [
    'Colectare dataset de imagini reale (poze de telefon cu exerciții) pentru îmbunătățirea scanner-ului',
    'Implementare Retrieval-Augmented Generation (RAG) cu baza de exerciții rezolvate',
    'Antrenare model mai mare (7B+) pe infrastructură GPU dedicată',
    'Adăugare suport pentru diagrame geometrice și grafice de funcții',
    'Evaluare sistematică pe subiecte BAC din anii anteriori (benchmark standardizat)',
    'Implementare feedback loop — modelul învață din greșelile raportate de utilizatori',
]
for f in future:
    add_bullet(f)

# Save
out = os.path.expanduser('~/Desktop/SmartBAC_AI_Capitol.docx')
doc.save(out)
print(f'Saved: {out}')
