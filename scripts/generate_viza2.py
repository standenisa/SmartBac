"""Generate Word — Viza 2: Arhitectură, BD, Protocoale, Fluxuri."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0x1B, 0x1B, 0x2F)
    hs.font.size = Pt(16 - (level-1)*2)
    hs.font.bold = True


def txt(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'})
    shading.append(shd)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()


# ═══════════════════════════════════════
# TITLE
# ═══════════════════════════════════════

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SmartBAC'); run.font.size = Pt(28); run.bold = True; run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x2F)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Viza 2 — Arhitectură, Baza de Date,\nProtocoale de Comunicație, Descrierea Fluxurilor')
run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nAplicație mobilă de pregătire pentru Bacalaureat\nla Matematică cu Inteligență Artificială')
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x8A)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n2026'); run.font.size = Pt(14)
doc.add_page_break()

# ═══════════════════════════════════════
# CUPRINS
# ═══════════════════════════════════════

doc.add_heading('Cuprins', level=1)
toc = [
    ('1.', 'Arhitectura generală a sistemului'),
    ('2.', 'Stack tehnologic'),
    ('3.', 'Structura bazei de date (MongoDB)'),
    ('4.', 'Protocoale de comunicație (REST API)'),
    ('5.', 'Endpointuri API — specificație completă'),
    ('6.', 'Fluxul de autentificare'),
    ('7.', 'Fluxul de rezolvare exerciții (Adaptive Learning)'),
    ('8.', 'Fluxul Chat AI (pipeline multi-model)'),
    ('9.', 'Fluxul Scanner (OCR + Solve)'),
    ('10.', 'Fluxul de gamificare'),
    ('11.', 'Diagrama de secvență — interacțiune utilizator'),
]
for num, title in toc:
    p = doc.add_paragraph(f'{num} {title}')
doc.add_page_break()

# ═══════════════════════════════════════
# 1. ARHITECTURA
# ═══════════════════════════════════════

doc.add_heading('1. Arhitectura generală a sistemului', level=1)

txt('SmartBAC urmează o arhitectură client-server cu trei componente principale: '
    'aplicația mobilă (frontend), serverul API (backend) și serviciile AI externe (Kaggle).')

doc.add_heading('1.1 Diagramă de componente', level=2)

code(
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                    FRONTEND (React Native / Expo)               │\n'
    '│  ┌────────┐ ┌──────────┐ ┌────────┐ ┌─────────┐ ┌──────────┐ │\n'
    '│  │ Acasă  │ │Exerciții │ │  Chat  │ │ Scanner │ │Predicție │ │\n'
    '│  └───┬────┘ └────┬─────┘ └───┬────┘ └────┬────┘ └────┬─────┘ │\n'
    '│      │           │           │            │           │        │\n'
    '│      └───────────┴─────┬─────┴────────────┴───────────┘        │\n'
    '│                        │ HTTP REST (JSON)                      │\n'
    '└────────────────────────┼────────────────────────────────────────┘\n'
    '                         │\n'
    '                         ▼\n'
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                    BACKEND (FastAPI + Python)                   │\n'
    '│  ┌──────────┐ ┌──────────┐ ┌───────────┐ ┌──────────────────┐ │\n'
    '│  │  Routers │ │ Services │ │  Models   │ │  AI Integration  │ │\n'
    '│  │ auth     │ │ adaptive │ │ user      │ │ math_tutor       │ │\n'
    '│  │ exercises│ │ gamific. │ │ exercise  │ │ math_explainer   │ │\n'
    '│  │ chat     │ │ leagues  │ │ attempt   │ │ recommender      │ │\n'
    '│  │ scanner  │ │ recomm.  │ │ achieve.  │ │ grade predictor  │ │\n'
    '│  │ stats    │ │ math_tut.│ │ exam      │ │ BPE tokenizer    │ │\n'
    '│  │ ml       │ │          │ │           │ │                  │ │\n'
    '│  └────┬─────┘ └──────────┘ └───────────┘ └────────┬─────────┘ │\n'
    '│       │                                           │           │\n'
    '└───────┼───────────────────────────────────────────┼───────────┘\n'
    '        │                                           │\n'
    '        ▼                                           ▼\n'
    '┌───────────────┐                    ┌──────────────────────────┐\n'
    '│   MongoDB     │                    │   Kaggle (via ngrok)     │\n'
    '│ bac_prep_ai   │                    │ ┌──────────────────────┐ │\n'
    '│               │                    │ │ DeepSeek-R1 (/ask)   │ │\n'
    '│ Collections:  │                    │ │ Chat + Solver        │ │\n'
    '│ users         │                    │ ├──────────────────────┤ │\n'
    '│ exercises     │                    │ │ Qwen VL (/transcrie) │ │\n'
    '│ attempts      │                    │ │ Scanner OCR          │ │\n'
    '│ achievements  │                    │ └──────────────────────┘ │\n'
    '│ leagues       │                    └──────────────────────────┘\n'
    '│ chat_history  │\n'
    '│ ...           │\n'
    '└───────────────┘'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 2. STACK TEHNOLOGIC
# ═══════════════════════════════════════

doc.add_heading('2. Stack tehnologic', level=1)

table(['Nivel', 'Tehnologie', 'Versiune', 'Rol'], [
    ['Frontend', 'React Native (Expo)', '54.x', 'Aplicație mobilă cross-platform'],
    ['Frontend', 'TypeScript', '5.9', 'Tipizare statică'],
    ['Frontend', 'Expo Router', '6.x', 'Navigare file-based'],
    ['Frontend', 'React Native Reanimated', '4.x', 'Animații native'],
    ['Frontend', 'KaTeX', '0.16', 'Randare formule matematice'],
    ['Backend', 'FastAPI', '3.0', 'Server REST API async'],
    ['Backend', 'Python', '3.14', 'Limbaj server'],
    ['Backend', 'PyMongo', '-', 'Driver MongoDB'],
    ['Backend', 'PyJWT', '-', 'Autentificare JWT'],
    ['Baza de date', 'MongoDB', '7.x', 'Stocare documente NoSQL'],
    ['AI — Chat', 'DeepSeek-R1-Distill-Qwen-14B', '14B params', 'Rezolvare exerciții (QLoRA)'],
    ['AI — Scanner', 'Qwen2.5-VL-3B', '3B params', 'OCR din imagini'],
    ['AI — Local', 'Rule-based solver', '-', 'Rezolvare offline'],
    ['AI — Predicție', 'Scikit-learn (Ensemble)', '-', 'Predicție notă BAC'],
    ['Infrastructură', 'Kaggle Notebooks', 'T4 GPU', 'Servire modele AI'],
    ['Tunel', 'ngrok', 'v3', 'Expunere server Kaggle'],
])

doc.add_page_break()

# ═══════════════════════════════════════
# 3. BAZA DE DATE
# ═══════════════════════════════════════

doc.add_heading('3. Structura bazei de date (MongoDB)', level=1)

txt('Baza de date utilizează MongoDB, o bază de date NoSQL orientată pe documente. '
    'Alegerea MongoDB se justifică prin: flexibilitatea schemei (exercițiile au câmpuri variabile), '
    'performanța la citiri frecvente și compatibilitatea nativă cu JSON.')

doc.add_heading('3.1 Colecția users', level=2)
table(['Câmp', 'Tip', 'Constrângeri', 'Descriere'], [
    ['_id', 'int', 'PK, auto-increment', 'Identificator unic'],
    ['email', 'string', 'unique, indexed', 'Adresa de email'],
    ['username', 'string', 'unique, indexed', 'Nume de utilizator'],
    ['password_hash', 'string', 'required', 'Parolă hashurită (PBKDF2)'],
    ['profile', 'string', 'enum: M1-M4', 'Profil BAC (Mate-Info, Științe, etc.)'],
    ['xp', 'int', 'default: 0', 'Puncte de experiență acumulate'],
    ['level', 'int', 'default: 1', 'Nivel curent (1-10)'],
    ['current_streak', 'int', 'default: 0', 'Serie zilnică curentă'],
    ['best_streak', 'int', 'default: 0', 'Cea mai lungă serie'],
    ['streak_freezes', 'int', 'default: 0', 'Protecții streak disponibile'],
    ['last_activity', 'datetime', '-', 'Ultima activitate'],
    ['created_at', 'datetime', 'auto', 'Data creării contului'],
])

doc.add_heading('3.2 Colecția exercises', level=2)
table(['Câmp', 'Tip', 'Index', 'Descriere'], [
    ['_id', 'int', 'PK', 'Identificator unic'],
    ['question', 'string', '-', 'Enunțul exercițiului'],
    ['answer', 'string', '-', 'Răspunsul corect'],
    ['difficulty', 'int', 'indexed', 'Nivel dificultate (1-5)'],
    ['topic', 'string', 'indexed', 'Categoria (ecuatie, derivata, etc.)'],
    ['subject', 'int', 'indexed', 'Subiect BAC (1, 2 sau 3)'],
    ['profile', 'string', 'indexed', 'Profil: M1, M2, M3, M4, BOTH'],
    ['exercise_type', 'string', 'indexed', 'Tipul exercițiului'],
    ['points', 'int', '-', 'XP acordat la rezolvare corectă'],
    ['solution_steps', 'array', '-', 'Pași de rezolvare'],
    ['hints', 'array', '-', 'Indicii progresive (3 nivele)'],
    ['latex', 'string', '-', 'Versiune LaTeX a enunțului'],
    ['year', 'int', '-', 'Anul subiectului BAC (dacă e cazul)'],
    ['session', 'string', '-', 'Sesiunea (Iunie, August)'],
])

doc.add_heading('3.3 Colecția attempts', level=2)
table(['Câmp', 'Tip', 'Index', 'Descriere'], [
    ['_id', 'int', 'PK', 'Identificator unic'],
    ['user_id', 'int', 'indexed', 'FK → users._id'],
    ['exercise_id', 'int', 'indexed', 'FK → exercises._id'],
    ['user_answer', 'string', '-', 'Răspunsul elevului'],
    ['is_correct', 'boolean', '-', 'Corectitudine'],
    ['time_spent', 'int', '-', 'Timp (secunde)'],
    ['created_at', 'datetime', 'indexed', 'Timestamp'],
])

doc.add_heading('3.4 Colecția user_exercise_history (Spaced Repetition)', level=2)
table(['Câmp', 'Tip', 'Descriere'], [
    ['user_id', 'int', 'FK → users'],
    ['exercise_id', 'int', 'FK → exercises'],
    ['quality', 'int', 'Calitate răspuns SM-2 (0-5)'],
    ['ease', 'float', 'Factor ușurință SM-2 (default 2.5)'],
    ['interval', 'int', 'Interval repetiție (zile)'],
    ['next_review', 'datetime', 'Data următoarei repetiții'],
    ['attempt_count', 'int', 'Număr încercări'],
])

doc.add_heading('3.5 Colecția leagues', level=2)
table(['Câmp', 'Tip', 'Descriere'], [
    ['user_id', 'int', 'FK → users'],
    ['week_start', 'string', 'Identificator săptămână (ISO Monday)'],
    ['league', 'string', 'Tier: Bronz, Argint, Aur, Diamant, Legenda'],
    ['weekly_xp', 'int', 'XP acumulat în săptămâna curentă'],
])

doc.add_heading('3.6 Alte colecții', level=2)
bullet('user_achievements — realizări deblocate (user_id, achievement_id, unlocked_at)')
bullet('exam_simulations — simulări examen (scoruri pe subiecte, timp, completare)')
bullet('daily_challenge_attempts — provocarea zilnică (user_id, date, is_correct)')
bullet('chat_history — istoric conversații chat (user_id, message, timestamp)')
bullet('counters — secvențe auto-increment pentru ID-uri')

doc.add_heading('3.7 Diagrama relațiilor', level=2)
code(
    'users ──────┐\n'
    '  │         │\n'
    '  │    ┌────┴────┐\n'
    '  ├───→│ attempts │←──── exercises\n'
    '  │    └─────────┘         │\n'
    '  │                        │\n'
    '  ├───→ user_achievements  │\n'
    '  │                        │\n'
    '  ├───→ user_exercise_history ←─┘\n'
    '  │     (spaced repetition)\n'
    '  │\n'
    '  ├───→ leagues\n'
    '  │\n'
    '  ├───→ exam_simulations\n'
    '  │\n'
    '  ├───→ daily_challenge_attempts\n'
    '  │\n'
    '  └───→ chat_history'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. PROTOCOALE
# ═══════════════════════════════════════

doc.add_heading('4. Protocoale de comunicație', level=1)

doc.add_heading('4.1 REST API (Frontend ↔ Backend)', level=2)
txt('Comunicația dintre aplicația mobilă și server se realizează prin protocolul HTTP/HTTPS '
    'folosind arhitectura REST (Representational State Transfer). Toate datele sunt transmise '
    'în format JSON.')

table(['Proprietate', 'Valoare'], [
    ['Protocol', 'HTTP/1.1 (local) / HTTPS (producție)'],
    ['Format date', 'JSON (application/json)'],
    ['Autentificare', 'JWT Bearer Token (header Authorization)'],
    ['CORS', 'Permis de pe toate originile (development)'],
    ['Timeout', '30 secunde (default)'],
    ['Port backend', '5005'],
    ['Port frontend', '8081'],
])

doc.add_heading('4.2 ngrok Tunnel (Backend ↔ Kaggle)', level=2)
txt('Modelele AI rulează pe GPU-uri Kaggle și sunt expuse prin tuneluri ngrok. '
    'Backend-ul FastAPI trimite cereri HTTP POST către URL-urile ngrok.')

table(['Tunel', 'Endpoint', 'Model', 'Payload'], [
    ['AI_NGROK_URL', '/ask', 'DeepSeek-R1', '{"intrebare": "text"}'],
    ['VLM_NGROK_URL', '/transcrie', 'Qwen VL', '{"image_base64": "base64"}'],
])

doc.add_heading('4.3 Formatul mesajelor', level=2)

txt('Cerere tipică (frontend → backend):')
code(
    'POST /api/exercises/submit-answer\n'
    'Authorization: Bearer eyJhbGciOiJIUzI1NiIs...\n'
    'Content-Type: application/json\n'
    '\n'
    '{\n'
    '  "user_id": 1,\n'
    '  "exercise_id": 42,\n'
    '  "answer": "x = 3",\n'
    '  "time_spent": 45\n'
    '}'
)

txt('Răspuns tipic (backend → frontend):')
code(
    'HTTP 200 OK\n'
    'Content-Type: application/json\n'
    '\n'
    '{\n'
    '  "correct": true,\n'
    '  "correct_answer": "x = 3",\n'
    '  "xp_earned": 10,\n'
    '  "new_achievements": ["first_correct"],\n'
    '  "message": "Corect! +10 XP"\n'
    '}'
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. ENDPOINTURI
# ═══════════════════════════════════════

doc.add_heading('5. Endpointuri API — specificație completă', level=1)

doc.add_heading('5.1 Autentificare (/api/auth)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['POST', '/auth/register', 'Creare cont nou'],
    ['POST', '/auth/login', 'Autentificare cu email/parolă'],
    ['GET', '/auth/me', 'Profil utilizator curent'],
    ['PUT', '/auth/me', 'Actualizare profil/parolă'],
    ['POST', '/auth/refresh', 'Reînnoire token JWT'],
])

doc.add_heading('5.2 Exerciții (/api/exercises)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['GET', '/exercises', 'Filtrare exerciții (topic, difficulty, profile)'],
    ['GET', '/exercises/next', 'Următorul exercițiu adaptiv (SM-2)'],
    ['GET', '/exercises/quick-practice', 'Exerciții din topicuri slabe'],
    ['POST', '/exercises/submit-answer', 'Trimite răspuns + verificare'],
    ['GET', '/exercises/{id}/solution', 'Soluție completă pas cu pas'],
    ['GET', '/exercises/{id}/hints', 'Indicii progresive (3 nivele)'],
])

doc.add_heading('5.3 Chat AI (/api/chat)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['POST', '/chat', 'Trimite mesaj la chatbot AI'],
])

doc.add_heading('5.4 Scanner (/api/scanner)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['POST', '/scanner/ocr', 'OCR imagine → text (Qwen VL pe Kaggle)'],
    ['POST', '/scanner/solve-text', 'Rezolvare text extras (DeepSeek R1)'],
])

doc.add_heading('5.5 Statistici și predicție (/api/stats, /api/ml)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['GET', '/stats', 'Statistici de bază (acuratețe)'],
    ['GET', '/stats/detailed', 'Statistici pe subiecte'],
    ['GET', '/stats/activity', 'Heatmap activitate'],
    ['GET', '/ml/predict-grade', 'Predicție notă BAC (ML ensemble)'],
    ['GET', '/ml/insights', 'Recomandări AI personalizate'],
])

doc.add_heading('5.6 Gamificare (/api/gamification)', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['GET', '/gamification/stats', 'XP, nivel, streak, achievements'],
    ['GET', '/gamification/achievements', 'Lista realizări cu status'],
    ['POST', '/gamification/streak/freeze', 'Activare protecție streak'],
])

doc.add_heading('5.7 Alte endpointuri', level=2)
table(['Metodă', 'Endpoint', 'Descriere'], [
    ['GET', '/daily-challenge', 'Provocarea zilnică'],
    ['POST', '/daily-challenge/submit', 'Răspuns provocare zilnică'],
    ['GET', '/leagues', 'Liga curentă + clasament'],
    ['GET', '/recommender/exercises', 'Recomandări exerciții'],
    ['POST', '/solver/solve', 'Rezolvare cu model AI'],
])

doc.add_page_break()

# ═══════════════════════════════════════
# 6-11. FLUXURI
# ═══════════════════════════════════════

doc.add_heading('6. Fluxul de autentificare', level=1)
code(
    '┌──────────┐    POST /auth/register     ┌──────────┐     Insert      ┌──────────┐\n'
    '│ Frontend │ ─────────────────────────→ │ Backend  │ ─────────────→ │ MongoDB  │\n'
    '│          │    {email, password,       │ FastAPI  │  users coll.   │          │\n'
    '│          │     username, profile}     │          │                │          │\n'
    '│          │ ←───────────────────────── │          │ ←───────────── │          │\n'
    '└──────────┘    {token, user}           └──────────┘   user_doc     └──────────┘\n'
    '\n'
    'Flux login:\n'
    '1. Frontend trimite POST /auth/login cu email + password\n'
    '2. Backend caută user-ul în MongoDB după email\n'
    '3. Verifică parola cu PBKDF2 hash\n'
    '4. Generează JWT token (valid 24h, HS256)\n'
    '5. Returnează token + date user\n'
    '6. Frontend stochează token-ul în AsyncStorage\n'
    '7. Toate cererile ulterioare includ: Authorization: Bearer <token>'
)

doc.add_page_break()

doc.add_heading('7. Fluxul de rezolvare exerciții (Adaptive Learning)', level=1)
code(
    '┌──────────┐  GET /exercises/next  ┌──────────┐  SM-2 Algorithm  ┌──────────┐\n'
    '│ Frontend │ ───────────────────→ │ Adaptive │ ───────────────→ │ MongoDB  │\n'
    '│          │                      │ Learning │                  │          │\n'
    '│          │ ←─────────────────── │ Service  │ ←─────────────── │          │\n'
    '│          │   {exercise}         │          │  history query   │          │\n'
    '└──────────┘                      └──────────┘                  └──────────┘\n'
    '\n'
    'Algoritm selecție (prioritate):\n'
    '  1. Exerciții cu repetiție scadentă (SM-2 interval expirat)\n'
    '  2. Exerciții din topicuri slabe (acuratețe < 50%)\n'
    '  3. Exerciții din topicuri neexplorate\n'
    '  4. Exerciții cu dificultate progresivă\n'
    '  5. Exerciții aleatorii (fallback)\n'
    '\n'
    'La submit răspuns:\n'
    '  1. Compară cu answer din exercițiu\n'
    '  2. Înregistrează attempt + actualizează SM-2\n'
    '  3. Acordă XP (10 XP corect)\n'
    '  4. Actualizează streak\n'
    '  5. Verifică achievements\n'
    '  6. Returnează: correct?, xp_earned, new_achievements'
)

doc.add_page_break()

doc.add_heading('8. Fluxul Chat AI (pipeline multi-model)', level=1)
code(
    '┌──────────┐  POST /api/chat  ┌───────────────────────────────────────────┐\n'
    '│ Frontend │ ───────────────→ │              Backend FastAPI              │\n'
    '│          │  {message}       │                                           │\n'
    '│          │                  │  1. Intent Detection (regex)              │\n'
    '│          │                  │     → solve / concept / re_explain        │\n'
    '│          │                  │                                           │\n'
    '│          │                  │  2. Normalizare input                     │\n'
    '│          │                  │     ² → ^2, √ → sqrt, radical → √        │\n'
    '│          │                  │     Adaugă prefix: "Rezolvă ecuația: ..." │\n'
    '│          │                  │                                           │\n'
    '│          │                  │  3. Trimite la DeepSeek R1 (Kaggle/ngrok) │\n'
    '│          │                  │     POST AI_NGROK_URL/ask                 │\n'
    '│          │                  │     {"intrebare": "...normalizat..."}     │\n'
    '│          │                  │                                           │\n'
    '│          │                  │  4. Dacă Kaggle offline → Rule-based      │\n'
    '│          │                  │     math_tutor.solve() (offline)          │\n'
    '│          │                  │                                           │\n'
    '│          │                  │  5. Parsare răspuns                       │\n'
    '│          │                  │     → pași, metodă, răspuns, verificare   │\n'
    '│          │                  │                                           │\n'
    '│          │                  │  6. Filtru anti-halucinare                │\n'
    '│          │                  │     → elimină gibberish, repetiții        │\n'
    '│          │ ←─────────────── │                                           │\n'
    '│          │  {structured,    │  7. Return JSON structurat                │\n'
    '│          │   suggestions}   └───────────────────────────────────────────┘\n'
    '│          │\n'
    '│  8. Randare KaTeX (formule) + SolutionView (pași)\n'
    '└──────────┘'
)

doc.add_page_break()

doc.add_heading('9. Fluxul Scanner (OCR + Solve)', level=1)
code(
    '┌──────────┐                    ┌──────────┐                ┌────────────────┐\n'
    '│ Frontend │  POST /scanner/ocr │ Backend  │  POST /transcrie │ Kaggle       │\n'
    '│          │ ─────────────────→ │ FastAPI  │ ────────────────→ │ Qwen VL      │\n'
    '│ 1. Poză  │  (image file)      │          │  {image_base64}  │ (ngrok)      │\n'
    '│ cameră/  │                    │          │                  │              │\n'
    '│ galerie  │                    │          │ ←──────────────── │              │\n'
    '│          │ ←───────────────── │          │  {transcriere}   │              │\n'
    '│          │  {extracted_text}  │          │                  └──────────────┘\n'
    '│          │                    │          │\n'
    '│ 2. User  │                    │          │                ┌────────────────┐\n'
    '│ editează │                    │          │                │ Kaggle         │\n'
    '│ textul   │                    │          │ POST /ask      │ DeepSeek R1    │\n'
    '│          │ POST /solve-text   │          │ ──────────────→│ (ngrok)        │\n'
    '│ 3. Click │ ─────────────────→ │          │ {intrebare}    │                │\n'
    '│ "Rezolvă"│                    │          │                │                │\n'
    '│          │                    │          │ ←──────────────│                │\n'
    '│          │ ←───────────────── │          │ {raspuns}      │                │\n'
    '│          │  {structured sol.} │          │                └────────────────┘\n'
    '│          │                    │          │\n'
    '│ 4. Afișare soluție cu KaTeX   └──────────┘\n'
    '└──────────┘'
)

doc.add_page_break()

doc.add_heading('10. Fluxul de gamificare', level=1)
code(
    'La fiecare răspuns corect:\n'
    '\n'
    '┌──────────┐   submit_answer   ┌──────────────┐   check_achievements   ┌──────────┐\n'
    '│ Exercises│ ─────────────────→│ Gamification │ ─────────────────────→ │ MongoDB  │\n'
    '│ Router   │                   │   Service    │                        │          │\n'
    '│          │ ←─────────────────│              │ ←───────────────────── │          │\n'
    '└──────────┘  {xp, achievem.}  └──────────────┘   query achievements   └──────────┘\n'
    '\n'
    'Sistem XP:\n'
    '  • Exercițiu corect: +10 XP\n'
    '  • Provocare zilnică (prima încercare): +25 XP\n'
    '  • Hint nivel 2: -5 XP  |  Hint nivel 3: -10 XP\n'
    '\n'
    'Nivele (10):\n'
    '  Începător(0) → Novice(100) → Elev(250) → Student(500) → Avansat(1000)\n'
    '  → Expert(2000) → Master(4000) → Guru(7000) → Legendă(11000) → Campion(16000)\n'
    '\n'
    'Achievements (8):\n'
    '  first_correct(10XP), streak_3(25XP), streak_5(50XP), streak_10(100XP)\n'
    '  exercises_10(30XP), exercises_50(100XP), exercises_100(250XP)\n'
    '  accuracy_80(100XP) — 80% acuratețe pe 20+ exerciții\n'
    '\n'
    'Ligi (resetare săptămânală):\n'
    '  Bronz(0XP) → Argint(100XP) → Aur(300XP) → Diamant(600XP) → Legendă(1000XP)\n'
    '  Top 5 promovează, ultimii 3 retrogradează'
)

doc.add_page_break()

doc.add_heading('11. Diagrama de secvență — interacțiune utilizator complet', level=1)
code(
    'Utilizator          Frontend            Backend           MongoDB         Kaggle\n'
    '   │                   │                   │                 │              │\n'
    '   │── Deschide app ──→│                   │                 │              │\n'
    '   │                   │── GET /auth/me ──→│── Find user ──→│              │\n'
    '   │                   │←── user data ─────│←── user doc ───│              │\n'
    '   │                   │                   │                 │              │\n'
    '   │── Tab Exerciții ─→│                   │                 │              │\n'
    '   │                   │── GET /next ─────→│── SM-2 query ─→│              │\n'
    '   │                   │←── exercise ──────│←── exercise ───│              │\n'
    '   │                   │                   │                 │              │\n'
    '   │── Scrie răspuns ─→│                   │                 │              │\n'
    '   │── Submit ────────→│── POST /submit ──→│── Insert ─────→│              │\n'
    '   │                   │                   │── Update XP ──→│              │\n'
    '   │                   │                   │── Check achiev→│              │\n'
    '   │                   │←── {correct,xp} ──│                │              │\n'
    '   │←── Confetti! ─────│                   │                 │              │\n'
    '   │                   │                   │                 │              │\n'
    '   │── Tab Chat ──────→│                   │                 │              │\n'
    '   │── "rezolvă x²=9"→│── POST /chat ────→│                │              │\n'
    '   │                   │                   │── Normalize ───│              │\n'
    '   │                   │                   │── POST /ask ───│─────────────→│\n'
    '   │                   │                   │                │              │\n'
    '   │                   │                   │←── {raspuns} ──│←─────────────│\n'
    '   │                   │                   │── Parse steps ─│              │\n'
    '   │                   │←── {structured} ──│                │              │\n'
    '   │←── KaTeX render ──│                   │                │              │\n'
    '   │                   │                   │                 │              │\n'
    '   │── Tab Scanner ───→│                   │                 │              │\n'
    '   │── Foto exercițiu→│── POST /ocr ─────→│                │              │\n'
    '   │                   │                   │── /transcrie ──│─────────────→│\n'
    '   │                   │                   │←── {text} ─────│←─────────────│\n'
    '   │                   │←── extracted_text─│                │              │\n'
    '   │── "Rezolvă" ────→│── POST /solve ───→│── /ask ────────│─────────────→│\n'
    '   │                   │                   │←── {raspuns} ──│←─────────────│\n'
    '   │                   │←── {solution} ────│                │              │\n'
    '   │←── Soluție ───────│                   │                │              │\n'
    '   ▼                   ▼                   ▼                ▼              ▼'
)

# Save
out = os.path.expanduser('~/Desktop/SmartBAC_Viza2.docx')
doc.save(out)
print(f'Saved: {out}')
